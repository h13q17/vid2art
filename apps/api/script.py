"""
если хочешь запустить код, то сначала нужно по очереди
ввести в консоль следующие строчки:

pip install youtube-transcript-api
pip install openai
pip install python-docx
pip install opencv-python

это установит необходимые сторонние библиотеки
"""

from youtube_transcript_api import YouTubeTranscriptApi # либа для получения субтитров
import openai                                           # либа для chatGPT
import docx                                             # либа для docx
import cv2                                              # либа для скриншотов из видео
import json
import math
import os

# OpenAI API key
openai.api_key = "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"

data = ''

with open('text.txt', 'r') as file:
    data = file.read();

parsed_data = json.loads(data)
link = parsed_data["url"]

print(link)

def cut_link(link):     # функция обрезки ссылки
    return link[len(link)-11:]

if len(link) > 11:      # если ссылка полная - обрезаем
    link = cut_link(link)
    
URL_base = "https://youtu.be/" + link + "?t="   # используется для создания гиперссылок


video_path = "orig_vid.mp4"     # название скаченного видео для скриншотов

def capture_screenshot(video_path, timecode, i):        # функция создания скриншота 
    
    # Получаем путь к папке, в которой находится файл программы
    current_dir = os.path.dirname(os.path.abspath(__file__))

    # Открываем видеофайл
    video = cv2.VideoCapture(video_path)

    # Устанавливаем таймкод для создания скриншота
    video.set(cv2.CAP_PROP_POS_MSEC, timecode * 1000)

    # Считываем текущий кадр
    success, image = video.read()
    
    # Сохраняем скриншот в файл
    screenshot_path = os.path.join(current_dir, ("screenshot_" + link + "_" + str(i) + ".jpg"))
    cv2.imwrite(screenshot_path, image)

    # Закрываем видеофайл
    video.release()

def add_hyperlink(paragraph, url, text, color, underline):      # функция добавления гиперссылки

    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # установка цвета
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # подчеркивание
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def create_part(subs_txt, begining_of_part, num_of_part):

    # генерация заголовка
    messages = [    # здесь идет указание для chatGPT без его ответа
        {"role": "system",
         "content" : "Озаглавь текст в несколько слов"}
    ]

    messages.append({"role": "user", "content": subs_txt})      # добавление основной части к запросу к chatGPT

    completion = openai.ChatCompletion.create(  # генерим ответ
        model="gpt-3.5-turbo-16k",
        messages=messages,
        n = 1,
        temperature = 0
    )

    chat_response_title = completion.choices[0].message.content

    # генерация основного текста
    messages = [    # здесь идет указание для chatGPT без его ответа
        {"role": "system",
         "content" : "Перескажи текст в одном абзаце"}
    ]

    messages.append({"role": "user", "content": subs_txt})      # добавление основной части к запросу к chatGPT

    completion = openai.ChatCompletion.create(  # генерим ответ
        model="gpt-3.5-turbo-16k",
        messages=messages,
        n = 1,
        temperature = 0
    )

    chat_response = completion.choices[0].message.content

    # добавление таймкода, заголовка и текста в .docx файл
    new_paragraph = result_docx.add_paragraph()
    _ = add_hyperlink(new_paragraph, URL_base + str(begining_of_part), ('[ ' + str(begining_of_part//60) + ':' + str(begining_of_part%60) + ' ]'), "Blue", True)
    title_of_part = result_docx.add_paragraph()
    runner = title_of_part.add_run(chat_response_title)
    runner.bold = True
    result_docx.add_paragraph(chat_response + '\n\n')


language = ['en']

transcript_list = YouTubeTranscriptApi.list_transcripts(link)   # если найдет среди субтитров
for transcript in transcript_list:                              # русский - значит статья будет
    if transcript.language_code == 'ru':                        # на нем, иначе англ
        language = ['ru']
        break

result_docx = docx.Document()
# .docx файл результата

# добавление пункта "Основано на:"
result_docx.add_paragraph("Основано на:")
new_paragraph = result_docx.add_paragraph()
_ = add_hyperlink(new_paragraph, "https://www.youtube.com/watch?v="+link, "https://www.youtube.com/watch?v="+link, "Blue", True)
result_docx.add_paragraph('\n')


subs_srt = YouTubeTranscriptApi.get_transcript(link, languages=language)
# переменная хранящая субтитры в формате srt

subs_txt = ''
# переменная хранящая субтитры в формате txt
for el in subs_srt:
    subs_txt += el['text']
    if (len(subs_txt) > 50000):
        break

# вступление-описание статьи
messages = [    # здесь идет указание для chatGPT без его ответа
    {"role": "system",
     "content" : "Расскажи про что этот текст в 2 предложениях"}
]

messages.append({"role": "user", "content": subs_txt})      # добавление основной части к запросу к chatGPT

completion = openai.ChatCompletion.create(  # генерим ответ
    model="gpt-3.5-turbo-16k",
    messages=messages,
    n = 1,
    temperature = 0
)

chat_response = completion.choices[0].message.content

result_docx.add_paragraph(chat_response + '\n\n')
    
video_duration = subs_srt[-1]['start']
# длина видео

# количество частей (можно/нужно редачить)
amount_of_parts = 5         # < 10 минут - 05 частей   
if video_duration > 600:    # > 10 минут - 10 частей
    amount_of_parts = 10    # > 01 часа  - 20 частей
if video_duration > 3600:
    amount_of_parts = 20

part_duration = math.floor(video_duration / amount_of_parts)
# длина одной части

subs_txt = ''                                            # переменная хранящая субтитры в формате txt
num_of_part = 1                                          # номер обрабатывающейся части
begining_of_part = int(math.floor(subs_srt[0]['start'])) # время начала части
for srt_el in subs_srt:
    
    subs_txt += srt_el['text']  # собираем все субтитры части в subs_txt
    
    if begining_of_part == -1:  # сохраняем время начала части
        begining_of_part = int(math.floor(srt_el['start']))
        
    if srt_el['start'] > part_duration * num_of_part:

        create_part(subs_txt, begining_of_part, num_of_part)    # создаем часть
        
        subs_txt = ''
        num_of_part += 1
        begining_of_part = -1
if subs_txt != '' :     # обработка последней "обрезанной" части
    create_part(subs_txt, begining_of_part, num_of_part)

#cохранение файла
result_file_name = "result_" + link + ".docx"
result_docx.save(result_file_name)
